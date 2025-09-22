import os
import sys
import argparse
import shutil
import time
import requests
import json
from pathlib import Path
from typing import List, Dict
from collections import deque

try:
    from git import Repo
except Exception:
    Repo = None

try:
    from docx import Document
    from docx.shared import Pt
except Exception:
    Document = None

# Supported code extensions (added .tsx for your repo)
CODE_EXTENSIONS = {
    ".py", ".js", ".jsx", ".ts", ".tsx", ".java", ".go", ".c", ".cpp", ".hpp",
    ".cs", ".rb", ".php", ".html", ".css", ".scss", ".yaml",
    ".yml", ".sh", ".rs", ".prisma"
}

# Excluded directories
EXCLUDE_DIRS = {
    "node_modules", "venv", ".venv", "env", ".env",
    "build", "dist", "__pycache__", ".git"
}

CHUNK_CHAR_SIZE = 8000
OLLAMA_API_URL = "http://localhost:11434/api/generate"


def parse_args():
    p = argparse.ArgumentParser(description="Generate developer documentation for a GitHub repo using Ollama.")
    p.add_argument("--repo", "-r", required=True)
    p.add_argument("--out", "-o", default="documentation.docx")
    p.add_argument("--tmp", "-t", default="./Code_Repository")
    p.add_argument("--model", default="gemma:2b") 
    p.add_argument("--keep", action="store_true")
    return p.parse_args()


def ensure_deps():
    missing = []
    if Repo is None:
        missing.append("gitpython (pip install GitPython)")
    if Document is None:
        missing.append("python-docx (pip install python-docx)")
    if missing:
        print("Missing dependencies:", file=sys.stderr)
        for m in missing:
            print("  -", m, file=sys.stderr)
        sys.exit(1)


def get_unique_repo_path(base_path: str) -> str:
    path = Path(base_path).resolve()
    if not path.exists():
        return str(path)

    counter = 1
    while True:
        new_path = path.parent / f"{path.name}_{counter}"
        if not new_path.exists():
            return str(new_path)
        counter += 1


def clone_repo(repo_url: str, dest: str) -> str:
    dest_path = Path(get_unique_repo_path(dest))
    Repo.clone_from(repo_url, str(dest_path))
    return str(dest_path)


# BFS-based file discovery
def find_code_files_bfs(root: str) -> List[str]:
    files = []
    queue = deque([root])

    while queue:
        current_dir = queue.popleft()
        try:
            for entry in os.scandir(current_dir):
                if entry.is_dir():
                    if entry.name not in EXCLUDE_DIRS:
                        queue.append(entry.path)
                else:
                    ext = os.path.splitext(entry.name)[1].lower()
                    if ext in CODE_EXTENSIONS:
                        files.append(entry.path)
        except PermissionError:
            continue  # skip folders we can’t open

    files.sort(key=lambda p: os.path.getsize(p))  # smallest first
    return files


def read_file(path: str) -> str:
    try:
        with open(path, "r", encoding="utf-8") as fh:
            return fh.read()
    except Exception:
        return ""


def ollama_generate(prompt: str, model: str) -> str:
    try:
        resp = requests.post(
            OLLAMA_API_URL,
            json={"model": model, "prompt": prompt},
            stream=True,
            timeout=300
        )
        resp.raise_for_status()
        output = ""
        for line in resp.iter_lines():
            if line:
                data = json.loads(line.decode("utf-8"))
                if "response" in data:
                    output += data["response"]
        return output.strip()
    except Exception as e:
        return f"[Error calling Ollama: {e}]"


def ollama_summarize_text_blocks(blocks: List[str], file_path: str, model: str = "gemma:2b") -> str:
    chunk_summaries = []
    for chunk in blocks:
        prompt = (
            "You are an expert software engineer and technical writer.\n"
            f"Analyze the following code excerpt from the file '{file_path}' and produce a clear, developer-friendly summary.\n\n"
            "For each excerpt, provide the following:\n"
            "1) One-sentence purpose.\n"
            "2) Key components.\n"
            "3) Functionality description.\n"
            "4) Dependencies.\n"
            "5) Edge cases or considerations.\n\n"
            f"Code excerpt:\n{chunk}\n\n"
        )
        text = ollama_generate(prompt, model)
        chunk_summaries.append(text)
        time.sleep(0.2)

    combined_prompt = (
        f"Combine these chunk-level summaries for file {file_path}:\n"
        "1) File summary\n2) Key components\n3) Actionable notes\n\n"
        "Chunk summaries:\n\n" + "\n\n---\n\n".join(chunk_summaries)
    )
    return ollama_generate(combined_prompt, model)


def chunk_text(text: str, size: int = CHUNK_CHAR_SIZE) -> List[str]:
    if not text:
        return []
    return [text[i:i+size] for i in range(0, len(text), size)]


def write_docx(output_path: str, repo_url: str, repo_root: str, deps_summary: str, file_summaries: Dict[str, str]):
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)
    doc.add_heading("Repository Documentation", level=0)
    doc.add_paragraph(f"Repository: {repo_url}")
    doc.add_paragraph(f"Scanned path: {repo_root}")
    doc.add_paragraph("Generated by: docgen.py (Ollama backend)")
    doc.add_heading("Dependencies & Top-level Files", level=1)
    doc.add_paragraph(deps_summary or "No dependency files were found.")
    doc.add_heading("File Summaries", level=1)
    for path, summary in file_summaries.items():
        doc.add_heading(path.replace(str(repo_root), "").lstrip(os.sep), level=2)
        doc.add_paragraph(summary)
    doc.add_page_break()
    doc.save(output_path)
    print(f"[+] Saved documentation to {output_path}")


def detect_top_level_info(root: str) -> str:
    candidates = ["requirements.txt", "Pipfile", "pyproject.toml", "package.json", "go.mod", "Gemfile"]
    found = []
    for name in candidates:
        p = Path(root) / name
        if p.exists():
            try:
                with p.open("r", encoding="utf-8") as fh:
                    sample = fh.read(2000)
                found.append(f"{name} (sample):\n{sample[:1000]}")
            except Exception:
                found.append(name)
    return "\n\n".join(found)


def main():
    args = parse_args()
    ensure_deps()
    repo_root = clone_repo(args.repo, args.tmp)
    try:
        deps_info = detect_top_level_info(repo_root)
        code_files = find_code_files_bfs(repo_root)  # swapped to BFS
        if not code_files:
            print("[!] No code files found. Exiting.")
            return
        file_summaries = {}
        for i, fpath in enumerate(code_files):
            rel = os.path.relpath(fpath, repo_root)
            print(f"[{i+1}/{len(code_files)}] Summarizing {rel}")
            content = read_file(fpath)
            if not content.strip():
                file_summaries[rel] = "[empty or binary file]"
                continue
            blocks = chunk_text(content, CHUNK_CHAR_SIZE)
            summ = ollama_summarize_text_blocks(blocks, rel, model=args.model)
            file_summaries[rel] = summ
        write_docx(args.out, args.repo, repo_root, deps_info, file_summaries)
    finally:
        if args.keep:
            print(f"[+] Keeping cloned repo at {repo_root}")
        else:
            try:
                shutil.rmtree(repo_root)
                print(f"[+] Removed temporary clone {repo_root}")
            except Exception:
                pass


if __name__ == "__main__":
    main()
